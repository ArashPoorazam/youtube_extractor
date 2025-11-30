import os
import logging
import docx 
from telegram import Update, ReplyKeyboardMarkup, KeyboardButton, ReplyKeyboardRemove
from telegram.ext import ContextTypes, CallbackContext

# Import Files
from youtube_extraction import YoutubeVideo
from database import add_or_update_user

logger = logging.getLogger(__name__)


# Download and send files with guaranteed cleanup
async def send_and_clean_file(update: Update, context: CallbackContext, download_func, file_type: str):
    """Handles download, sending, and required file cleanup."""
    
    # 1. Retrieve the link from user_data
    link = context.user_data.get('video_link')
    if not link:
        await update.message.reply_text("❌ لطفا اول لینک ویدیو را بفرستید.")
        return

    # 2. Initialize video object and download
    path = None 
    try:
        video = YoutubeVideo(link)
        await update.message.reply_text(f"⏳ کیفیت {file_type} لطفا منتظر بمانید، در حال دانلود...")
        
        path = download_func(video)
        caption = "📥 دانلود سریع" + " | " + "@Aroura"
        if path:
            if file_type == "Audio":
                await update.message.reply_audio(audio=path, caption=caption)
            elif file_type.startswith("Video"):
                await update.message.reply_video(video=path, caption=caption)
            
            await update.message.reply_text(f"کیفیت {file_type} با موفقت ارسال شد!")
        else:
            await update.message.reply_text(f"کیفیت {file_type} برای این ویدیو پیدا نشد...")

    except Exception as e:
        logger.error(f"Error during {file_type} processing: {e}", exc_info=True)
        await update.message.reply_text(f"خطایی هنگام پردازش رخ داد، دوباره تلاش کنید.")

    # 3. Ensure file deletion 
    finally:
        if path and os.path.exists(path):
            try:
                os.remove(path)
                logger.info(f"Cleaned up file: {path}")
            except OSError as e:
                logger.error(f"Error deleting file {path}: {e}")


### commands
async def start_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    user = update.effective_user
    
    # 1. Clear context data
    context.user_data.clear()
    
    # 2. Database interaction: Add/Update user info
    if user:
        add_or_update_user(
            user_id=user.id,
            first_name=user.first_name,
            last_name=user.last_name,
            username=user.username
        )
        logger.info(f"User {user.id} started the bot. user_data cleared and user recorded/updated.")
    
    # 3. Send the welcome message
    await context.bot.send_message(
        chat_id=update.effective_chat.id, 
        text="😁👋 سلام به ربات دانلود ویدیو یوتوب آرورا خوش آمدید.\n"
    )

async def help_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    message = (
        "😁👋 سلام به ربات دانلود ویدیو یوتوب آرورا خوش آمدید.\n"        
        "فقط کافیه ( لینک ) ویدیو یوتوب برای ربت بفرستید تا ویدیو یا صدا و یا زیر نویس اش رو از یوتوب دانلود کنید.\n\n"
        "🟢 گزینه ها:\n\n"
        "🎥 ویدیو - 144p - 360p - 720p- 1080p\n"
        "🈯 زیر نویس - روسی 🇷🇺 - انگلیسی 🇺🇸 (در قالب Word Document)\n"
        "🔊 صدا با کیفیت ترین حالت ممکنه\n"
    )
    await update.message.reply_text(message)

async def creator_command(update: Update, context: ContextTypes.DEFAULT_TYPE):
    creator_info = (
        "<b>Bot Creator Information:</b>\n\n"
        "👤 **Name:** [Arash Poorazam]\n"
        "🔗 **GitHub:** <a href='https://github.com/ArashPoorazam'>My GitHub Profile</a>\n"
        "📧 **Email:** <a href='arashpoorazam@gmail.com'>youremail@example.com</a>\n"
        "💼 **LinkedIn:** <a href='https://www.linkedin.com/in/arash-poorazam-b3a6a8292/'>My LinkedIn Profile</a>"
    )

    await context.bot.send_message(
        chat_id=update.effective_chat.id,
        text=creator_info,
        parse_mode='HTML' 
    )


### buttons
async def link_buttons(update: Update, context: CallbackContext, link: str):
    link = context.user_data.get('video_link')
    if not link:
        await update.message.reply_text("❌ لطفا اول لینک ویدیو را بفرستید.")
        return
    
    video = YoutubeVideo(link)

    keyboard = [
        [KeyboardButton("🎥 Video"), KeyboardButton("🔊 Audio")],
        [KeyboardButton("🈯 Subtitle")],
        [KeyboardButton("Go Back")]
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True, one_time_keyboard=True)
    await update.message.reply_text(text=f"{video.yt.title}\n\nلینک: {link}", reply_markup=reply_markup)
    await update.message.reply_text(text="چه کاری میتونم براتون انجام بدم؟ 😁", reply_markup=reply_markup)


async def video_q_buttons(update: Update, context: CallbackContext):
    keyboard = [
        [KeyboardButton("🎥 144 P"), KeyboardButton("🎥 360 P")],
        [KeyboardButton("🎥 720 P"), KeyboardButton("🎥 1080 P")],
        [KeyboardButton("Go Back")]
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True, one_time_keyboard=True)
    await update.message.reply_text(text="با چه کیفیتی میخواید دانلود کنید؟", reply_markup=reply_markup)


# subtitle
async def sub_choose(update: Update, context: CallbackContext):
    link = context.user_data.get('video_link')

    if not link:
        await update.message.reply_text("❌ لطفا اول لینک ویدیو را بفرستید.")
        return

    keyboard = [
        [KeyboardButton("🇺🇸 English"), KeyboardButton("🇷🇺 Russia")],
        [KeyboardButton("Go Back")]
    ]

    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True, one_time_keyboard=True)
    await update.message.reply_text(text="زیر نویس به چه زبانی باشد؟ (در قالب Word Document)", reply_markup=reply_markup)


async def send_subtitle_docx(update: Update, context: CallbackContext, lang_code: str):
    link = context.user_data.get('video_link')
    if not link:
        await update.message.reply_text("❌ لطفا اول لینک ویدیو را بفرستید.")
        return

    docx_path = None
    # Map language codes to display name and the appropriate retrieval function in YoutubeVideo
    lang_map = {
        'en': ('انگلیسی', YoutubeVideo.get_pure_subtitles_text), 
        'ru': ('روسی', YoutubeVideo.get_pure_subtitles_text)
    }
    
    if lang_code not in lang_map:
        await update.message.reply_text("زبان نامعتبر.")
        return
        
    lang_name, get_caption_func = lang_map[lang_code]

    try:
        video = YoutubeVideo(link)
        
        # Dynamically call the correct subtitle method, passing the required lang_code
        caption = get_caption_func(video, lang_code)
        
        if not caption:
            await update.message.reply_text(f"زیر نویسی برای زبان {lang_name} یافت نشد.")
            return

        # 1. Create the DOCX file
        video_title = video.yt.title

        safe_title = "".join(c for c in video_title if c.isalnum() or c in (' ', '_', '-')).strip()
        filename = f"{safe_title}_{lang_code}_subtitles" 
        
        await update.message.reply_text("⏳ در حال تولید فایل Word زیرنویس، لطفا منتظر بمانید...")
        
        # Use the new DOCX creation function
        docx_path = create_subtitle_docx(caption, filename) 
        
        if docx_path:
            # 2. Send the DOCX document
            with open(docx_path, 'rb') as docx_file:
                await update.message.reply_document(
                    document=docx_file,
                    filename=os.path.basename(docx_path), 
                    caption=f"📝 زیرنویس {lang_name} ویدیو در قالب Word Document: {video_title}"
                )
            await update.message.reply_text("فایل Word زیرنویس با موفقیت ارسال شد.")
        else:
            await update.message.reply_text("❌ خطایی هنگام تولید فایل Word رخ داد.")

    except Exception as e:
        logger.error(f"Error getting {lang_name} subtitles and sending DOCX: {e}", exc_info=True)
        await update.message.reply_text("خطایی هنگام پردازش رخ داد دوباره تلاش کنید.")

    # 3. Clean up the generated DOCX file
    finally:
        if docx_path and os.path.exists(docx_path):
            try:
                os.remove(docx_path)
                logger.info(f"Cleaned up DOCX file: {docx_path}")
            except OSError as e:
                logger.error(f"Error deleting DOCX file {docx_path}: {e}")


def create_subtitle_docx(text_content: str, filename: str) -> str:
    """Creates a .docx file containing the subtitle text content."""
    
    document = docx.Document()
    document.add_heading('YouTube Video Subtitles', 0)
    
    try:
        paragraph = text_content.replace(".", ".\n")
        document.add_paragraph(paragraph)
        base_name, _ = os.path.splitext(filename)
        docx_filename = f"{base_name}.docx"
        filepath = os.path.join("videos", docx_filename)

        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        document.save(filepath)
        
        logger.info(f"Successfully created DOCX file: {filepath}")
        return filepath
    
    except Exception as e:
        logger.error(f"Failed to create DOCX file: {e}")
        return None


# Go back
async def go_back(update: Update, context: CallbackContext): 
    keyboard = [
        [KeyboardButton("🎥 Video"), KeyboardButton("🔊 Audio")],
        [KeyboardButton("🈯 Subtitle")],
        [KeyboardButton("Go Back")]
    ]
    reply_markup = ReplyKeyboardMarkup(keyboard, resize_keyboard=True, one_time_keyboard=True)
    await update.message.reply_text("🏡 بازگشت.", reply_markup=ReplyKeyboardRemove())
    await update.message.reply_text(text="چه کاری میتونم براتون انجام بدم؟ 😁", reply_markup=reply_markup)


### chats
async def chat_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    text = update.message.text
    logger.debug(f"Handling unrecognised text: {text}")
    await update.message.reply_text(
        f"Aurora is coming soon...\nReceived: {text}"
    )


# --- Main Message Handler ---
async def handle_messages(update: Update, context: CallbackContext):
    text = update.message.text
    logger.debug(f"Received message: {text}")
    
    if text.startswith("https://youtu.be/") or text.startswith("https://youtube.com/"):
        # Store the link in user_data
        context.user_data['video_link'] = text 
        logger.info(f"New video link stored in user_data: {text}")
        await link_buttons(update, context, text)
        return

    match text:
        case "Go Back":
            await go_back(update, context) 
        case "🈯 Subtitle":
            await sub_choose(update, context)
        case "🎥 Video":
            await video_q_buttons(update, context)
        case "🔊 Audio":
            await send_and_clean_file(update, context, YoutubeVideo.download_audio, "Audio")
        case "🎥 144 P":
            await send_and_clean_file(update, context, YoutubeVideo.download_video_144, "Video 144p")
        case "🎥 360 P":
            await send_and_clean_file(update, context, YoutubeVideo.download_video_360, "Video 360p")
        case "🎥 720 P":
            await send_and_clean_file(update, context, YoutubeVideo.download_video_720, "Video 720p")
        case "🎥 1080 P":
            await send_and_clean_file(update, context, YoutubeVideo.download_video_1080, "Video 1080p")
        case "🇺🇸 English":
            # Call the new DOCX function
            await send_subtitle_docx(update, context, 'en')
        case "🇷🇺 Russia":
            # Call the new DOCX function
            await send_subtitle_docx(update, context, 'ru')
        case _:
            await chat_handler(update, context)


async def export_users(update: Update, context: ContextTypes.DEFAULT_TYPE):
    ADMIN_USERNAME = os.getenv("ADMIN_USERNAME")
    
    if update.effective_user.username != ADMIN_USERNAME:
        logger.info("Not Admin...!")
        return
    try:
        await update.message.reply_document(
            document=open('bot_users.db', 'rb'),
            filename='users_backup.db',
            caption="Here is the latest user database. 📂"
        )
        logger.info("Exported users...!")
    except Exception as e:
        await update.message.reply_text(f"Error sending database: {e}")


# Errors
async def error_handler(update: Update, context: ContextTypes.DEFAULT_TYPE):
    logger.error("Exception while handling an update:", exc_info=context.error)